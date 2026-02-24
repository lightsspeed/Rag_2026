from fastapi import APIRouter, UploadFile, File, HTTPException, Depends, BackgroundTasks, WebSocket, WebSocketDisconnect, Request
from fastapi.responses import FileResponse, HTMLResponse
from sqlalchemy.orm import Session
from typing import List, Optional
import shutil
import os
import uuid
import time
import math
import re
import hashlib
from datetime import datetime
from pydantic import BaseModel

from app.db.postgres import get_db, SessionLocal
from app.db import models
from app.services.retriever import retriever
from app.services.generator import generator
from app.services.cache import redis_cache
from app.services.ingestion import ingestion_service
from app.services.web_search import web_search_service
from app.core.limiter import limiter
from app.core.security import get_current_user, get_user_from_token_str

router = APIRouter()


# --- Helpers ---

def sigmoid(x: float) -> float:
    return 1 / (1 + math.exp(-x))


# --- Data Models ---

class FeedbackRequest(BaseModel):
    conversation_id: Optional[str] = None
    message_id: Optional[str] = None
    rating: Optional[str] = None        # "up" or "down"
    message_preview: Optional[str] = None
    query_preview: Optional[str] = None
    # legacy fields kept for backward compat
    query_id: Optional[int] = None
    score: Optional[int] = None


class TitleRequest(BaseModel):
    query: str
    conversation_id: Optional[str] = None


# --- Document Endpoints ---

@router.post("/documents/upload")
@limiter.limit("5/minute")
async def upload_document(
    request: Request,
    file: UploadFile = File(...),
    background_tasks: BackgroundTasks = BackgroundTasks(),
    db: Session = Depends(get_db),
):
    temp_dir = "uploads"
    os.makedirs(temp_dir, exist_ok=True)
    file_path = os.path.join(temp_dir, file.filename)

    with open(file_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)

    file.file.seek(0)
    file_hash = hashlib.sha256(file.file.read()).hexdigest()
    file.file.seek(0)

    existing_doc = db.query(models.Document).filter(models.Document.file_hash == file_hash).first()
    if existing_doc:
        return {
            "filename": existing_doc.filename,
            "status": existing_doc.status,
            "document_id": existing_doc.id,
        }

    db_doc = models.Document(
        filename=file.filename,
        file_hash=file_hash,
        status="processing",
    )
    db.add(db_doc)
    db.commit()
    db.refresh(db_doc)

    background_tasks.add_task(ingestion_service.process_document, file_path, file.filename, file_hash, db)

    return {
        "filename": db_doc.filename,
        "status": db_doc.status,
        "document_id": db_doc.id,
    }


@router.get("/documents")
def list_documents(
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_user),
):
    docs = db.query(models.Document).order_by(models.Document.upload_date.desc()).all()
    return [
        {
            "id": d.id,
            "filename": d.filename,
            "upload_date": d.upload_date.isoformat() if d.upload_date else "",
            "status": d.status,
            "chunk_count": d.chunk_count or 0,
        }
        for d in docs
    ]


@router.delete("/documents/{document_id}")
def delete_document(
    document_id: int,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_user),
):
    doc = db.query(models.Document).filter(models.Document.id == document_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    # Remove from filesystem
    file_path = os.path.join("uploads", doc.filename)
    if os.path.exists(file_path):
        os.remove(file_path)

    db.delete(doc)
    db.commit()
    return {"status": "deleted"}


@router.get("/documents/download/{filename}")
def download_document(
    filename: str,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_user),
):
    file_path = os.path.join("uploads", filename)
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="File not found")
    return FileResponse(file_path, filename=filename)


@router.get("/documents/preview/{filename}")
def preview_document(
    filename: str,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_user),
):
    file_path = os.path.join("uploads", filename)
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="File not found")

    ext = os.path.splitext(filename)[1].lower()

    # PDFs are returned as-is — browsers render them natively in an iframe
    if ext == ".pdf":
        return FileResponse(file_path, media_type="application/pdf")

    # For DOCX extract text via python-docx
    body_html = ""
    if ext == ".docx":
        try:
            import docx as _docx
            doc = _docx.Document(file_path)
            parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    import html as _html
                    parts.append(f"<p>{_html.escape(para.text)}</p>")
            for table in doc.tables:
                rows_html = ""
                for row in table.rows:
                    cells = "".join(f"<td style='border:1px solid #ddd;padding:6px'>{_html.escape(cell.text.strip())}</td>" for cell in row.cells)
                    rows_html += f"<tr>{cells}</tr>"
                parts.append(f"<table style='border-collapse:collapse;width:100%;margin:1em 0'>{rows_html}</table>")
            body_html = "\n".join(parts) if parts else "<p><em>No readable content found.</em></p>"
        except Exception as e:
            body_html = f"<p><em>Could not parse document: {e}</em></p>"
    else:
        # TXT / MD / other text files
        try:
            import html as _html
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read()
            body_html = f"<pre style='white-space:pre-wrap;word-break:break-word'>{_html.escape(text)}</pre>"
        except Exception as e:
            body_html = f"<p><em>Could not read file: {e}</em></p>"

    page = f"""<!DOCTYPE html>
<html lang="en">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width,initial-scale=1">
  <style>
    body{{font-family:Arial,sans-serif;padding:2rem;line-height:1.7;color:#222;max-width:960px;margin:0 auto;font-size:14px}}
    p{{margin:0 0 .75em}}
    pre{{font-size:13px;font-family:monospace}}
    table{{margin:1em 0}}
    td{{vertical-align:top}}
  </style>
</head>
<body>{body_html}</body>
</html>"""
    return HTMLResponse(content=page)


# --- WebSocket Chat ---

@router.websocket("/ws/chat")
async def websocket_endpoint(websocket: WebSocket):
    print("New WebSocket connection attempt...")
    await websocket.accept()
    print("WebSocket connection accepted.")

    # Extract token from query params for user identification
    token = websocket.query_params.get("token")

    try:
        while True:
            try:
                data = await websocket.receive_json()
                query = data.get("query")
                session_id = data.get("session_id")
                user_id = data.get("user_id", "anonymous")
                user_name = data.get("user_name", "User")

                if not query:
                    continue

                # Resolve user from token if available
                db = SessionLocal()
                try:
                    resolved_user_id = user_id
                    if token:
                        db_user = get_user_from_token_str(token, db)
                        if db_user:
                            resolved_user_id = db_user.id

                    # Ensure Conversation record exists for this session
                    if session_id:
                        conv = db.query(models.Conversation).filter(
                            models.Conversation.id == session_id
                        ).first()
                        if not conv:
                            conv = models.Conversation(
                                id=session_id,
                                user_id=resolved_user_id if resolved_user_id != "anonymous" else None,
                                title=None,
                            )
                            db.add(conv)
                            db.commit()

                    # 0. Context Management
                    history = redis_cache.get_session(session_id, resolved_user_id) or []

                    # 0.5 Intent Check
                    greeting_pattern = r"^(hi|hello|hey|greetings|good\s*(morning|afternoon|evening)|thanks|thank\s*you|bye|goodbye|hii+)(\s.*)?$"
                    is_obvious_greeting = re.match(greeting_pattern, query, re.IGNORECASE)

                    if is_obvious_greeting:
                        intent = {"category": "GREETING", "reply": "Hello! How can I assist you with your IT issues today?"}
                    else:
                        intent = await generator.check_intent(query)

                    if intent.get("category") in ["GREETING", "NONSENSE"]:
                        direct_reply = intent.get("reply", "How can I help you?")
                        turn_id = str(uuid.uuid4())
                        for char in direct_reply:
                            await websocket.send_json({"type": "token", "content": char})
                        await websocket.send_json({"type": "complete", "message_id": turn_id})

                        # Save turn
                        if session_id:
                            turn = models.ConversationTurn(
                                id=turn_id,
                                conversation_id=session_id,
                                query=query,
                                response=direct_reply,
                                sources_json=[],
                            )
                            db.add(turn)
                            db.commit()

                        history.append({"role": "user", "content": query})
                        history.append({"role": "assistant", "content": direct_reply})
                        redis_cache.update_session(session_id, resolved_user_id, history)
                        continue

                    # Standalorize Query
                    effective_query = await generator.standalorize_query(history, query)

                    # 1. Retrieve
                    chunks = await retriever.retrieve(effective_query)

                    is_web_search = False
                    if not chunks:
                        is_web_search = True
                        await websocket.send_json({"type": "status", "content": "Searching the web for more information..."})
                        chunks = await web_search_service.search(effective_query)

                    # Send Sources (mapped to new frontend format)
                    sources = []
                    for i, chunk in enumerate(chunks):
                        raw_score = chunk.get("score", 0)
                        confidence = round(sigmoid(raw_score), 2)
                        metadata = chunk.get("metadata", {})
                        is_web = bool(metadata.get("is_web", False))
                        if is_web:
                            doc_name = metadata.get("title") or "Web Result"
                        else:
                            doc_name = metadata.get("filename") or metadata.get("source") or "Document"
                        sources.append({
                            "id": chunk.get("id") or f"chunk-{i}",
                            "documentName": doc_name,
                            "excerpt": chunk["text"][:300],
                            "confidence": confidence,
                            "isWeb": is_web,
                            "url": metadata.get("url", "") if is_web else "",
                        })
                    await websocket.send_json({"type": "sources", "sources": sources})

                    # 2. Generate (Stream)
                    full_ai_response = ""
                    async for token_text in generator.generate_stream(effective_query, chunks):
                        await websocket.send_json({"type": "token", "content": token_text})
                        full_ai_response += token_text

                    turn_id = str(uuid.uuid4())
                    await websocket.send_json({"type": "complete", "message_id": turn_id})

                    # 3. Save History
                    history.append({"role": "user", "content": query})
                    history.append({"role": "assistant", "content": full_ai_response})
                    redis_cache.update_session(session_id, resolved_user_id, history)

                    # 4. Save Conversation Turn
                    if session_id:
                        turn = models.ConversationTurn(
                            id=turn_id,
                            conversation_id=session_id,
                            query=query,
                            response=full_ai_response,
                            sources_json=sources,
                        )
                        db.add(turn)

                        # Update conversation timestamp
                        conv = db.query(models.Conversation).filter(
                            models.Conversation.id == session_id
                        ).first()
                        if conv:
                            conv.updated_at = datetime.utcnow()
                        db.commit()

                    # 5. Log
                    log = models.QueryLog(
                        user_id=resolved_user_id,
                        query_text=effective_query,
                        retrieved_chunks=len(chunks),
                        response_time_ms=0,
                    )
                    db.add(log)
                    db.commit()

                finally:
                    db.close()

            except WebSocketDisconnect:
                raise
            except Exception as e:
                print(f"Error handling WebSocket query: {e}")
                import traceback; traceback.print_exc()
                try:
                    await websocket.send_json({"type": "error", "content": f"An error occurred: {str(e)}"})
                except Exception:
                    pass
                continue

    except WebSocketDisconnect:
        print("Client disconnected")


# --- Feedback ---

@router.post("/feedback")
def submit_feedback(
    feedback: FeedbackRequest,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_user),
):
    fb = models.Feedback(
        conversation_id=feedback.conversation_id,
        message_id=feedback.message_id,
        user_id=current_user.id,
        rating=feedback.rating,
        message_preview=feedback.message_preview,
        query_preview=feedback.query_preview,
    )
    db.add(fb)

    # Also update ConversationTurn feedback if message_id exists
    if feedback.message_id:
        turn = db.query(models.ConversationTurn).filter(
            models.ConversationTurn.id == feedback.message_id
        ).first()
        if turn:
            turn.feedback = feedback.rating

    db.commit()
    db.refresh(fb)
    return {"status": "ok", "id": fb.id}


# --- Chat Title ---

@router.post("/chat/title")
@limiter.limit("20/minute")
async def generate_chat_title(request: Request, request_body: TitleRequest, db: Session = Depends(get_db)):
    title = generator.generate_title(request_body.query)

    # Update conversation title if conversation_id provided
    if request_body.conversation_id and title:
        conv = db.query(models.Conversation).filter(
            models.Conversation.id == request_body.conversation_id
        ).first()
        if conv and not conv.title:
            conv.title = title
            db.commit()

    return {"title": title}
