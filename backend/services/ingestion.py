import os
import uuid
import time
import logging
from sqlalchemy.orm import Session
from backend.db import models
from backend.services.chunker import chunker
from backend.services.embedder import embedder
from backend.db.chroma import get_collection

logger = logging.getLogger(__name__)

class IngestionService:
    def process_all_in_dir(self, directory: str, db: Session):
        """Scans directory and processes all files."""
        if not os.path.exists(directory):
            logger.warning(f"Directory {directory} does not exist. Skipping startup scan.")
            return

        files = [f for f in os.listdir(directory) if os.path.isfile(os.path.join(directory, f))]
        if not files:
            logger.info(f"Folder {directory} is empty.")
            return

        logger.info(f"Found {len(files)} files in {directory}. Starting batch processing...")

        for filename in files:
            file_path = os.path.join(directory, filename)
            file_hash = str(uuid.uuid5(uuid.NAMESPACE_DNS, filename)) 
            
            existing = db.query(models.Document).filter(models.Document.filename == filename).first()
            if existing and existing.status == "completed":
               continue

            try:
                if not existing:
                    db_doc = models.Document(
                        filename=filename,
                        file_hash=file_hash,
                        status="processing"
                    )
                    db.add(db_doc)
                    db.commit()
                
                self.process_document(file_path, filename, file_hash, db)
            except Exception as e:
                logger.error(f"Failed to initiate processing for {filename}: {e}")

    def process_document(self, file_path: str, filename: str, file_hash: str, db: Session, progress_callback=None):
        start_time = time.time()
        logger.info(f"START PROCESSING: {filename}")
        
        def report(status, progress=0, details=None):
            if progress_callback:
                try:
                    progress_callback({
                        "type": "ingestion_progress",
                        "filename": filename,
                        "status": status,
                        "progress": progress,
                        "details": details
                    })
                except Exception as e:
                    logger.error(f"Failed to report progress: {e}")

        report("reading", 10, "Extracting text from file...")
        content = ""
        ext = os.path.splitext(filename)[1].lower()
        
        try:
            if ext == ".pdf":
                try:
                    import pdfplumber
                    with pdfplumber.open(file_path) as pdf:
                        for page in pdf.pages:
                            text_content = page.extract_text(layout=True)
                            if text_content:
                                content += text_content + "\n"
                            tables = page.extract_tables()
                            for table in tables:
                                table_str = "\n[Table Start]\n"
                                for row in table:
                                    clean_row = [str(cell) if cell is not None else "" for cell in row]
                                    table_str += " | ".join(clean_row) + "\n"
                                table_str += "[Table End]\n"
                                content += table_str
                except Exception as e:
                    logger.error(f"pdfplumber error: {e}")
                    import pypdf
                    reader = pypdf.PdfReader(file_path)
                    for page in reader.pages:
                        extracted = page.extract_text()
                        if extracted:
                            content += extracted + "\n"
            elif ext == ".docx":
                import docx
                doc = docx.Document(file_path)
                for para in doc.paragraphs:
                    if para.text.strip():
                        content += para.text + "\n"
                for table in doc.tables:
                    table_str = "\n[Table Start]\n"
                    for row in table.rows:
                        clean_row = [cell.text.strip() for cell in row.cells]
                        table_str += " | ".join(clean_row) + "\n"
                    table_str += "[Table End]\n"
                    content += table_str
            else:
                with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                    content = f.read()

            if not content.strip():
                logger.warning(f"No content extracted from {filename}")
                self._mark_failed(db, file_hash)
                report("failed", 100, "Extraction returned empty content.")
                return

            # 2. Chunk
            report("chunking", 30, "Dividing content into chunks...")
            metadata = {
                'source': filename, 
                'filename': filename,
                'upload_time': time.time(), 
                'file_hash': file_hash
            }
            chunks = chunker.chunk_text(content, metadata)
            chunk_count = len(chunks)

            if not chunks:
                self._mark_failed(db, file_hash)
                report("failed", 100, "No chunks generated.")
                return

            # 3. Embed
            report("embedding", 50, f"Generating embeddings for {chunk_count} chunks...")
            texts = [chunk['text'] for chunk in chunks]
            embeddings = embedder.embed_batch(texts) 

            # 4. Store in Chroma and SQLite
            report("storing", 80, f"Saving {chunk_count} vectors to database...")
            collection = get_collection()
            
            doc = db.query(models.Document).filter(models.Document.file_hash == file_hash).first()
            if not doc:
                self._mark_failed(db, file_hash)
                report("failed", 100, "Document record lost.")
                return

            doc.content = content
            for i, chunk in enumerate(chunks):
                vector_id = str(uuid.uuid4())
                collection.add(
                    documents=[chunk['text']],
                    embeddings=[embeddings[i]],
                    metadatas=[chunk['metadata']],
                    ids=[vector_id]
                )
                db_chunk = models.Chunk(
                    document_id=doc.id,
                    vector_id=vector_id,
                    content=chunk['text'],
                    summary=chunk['metadata'].get('summary', ''),
                    keywords=chunk['metadata'].get('keywords', []),
                    questions=chunk['metadata'].get('questions', [])
                )
                db.add(db_chunk)

            doc.status = "completed"
            doc.chunk_count = len(chunks)
            db.commit()
            
            elapsed_time = time.time() - start_time
            report("completed", 100, f"Processed {chunk_count} chunks in {elapsed_time:.1f}s")
                
        except Exception as e:
            logger.error(f"Error processing {filename}: {e}")
            self._mark_failed(db, file_hash)
            report("failed", 100, f"Error: {str(e)}")

    def _mark_failed(self, db, file_hash):
        doc = db.query(models.Document).filter(models.Document.file_hash == file_hash).first()
        if doc:
            doc.status = "failed"
            db.commit()

ingestion_service = IngestionService()
